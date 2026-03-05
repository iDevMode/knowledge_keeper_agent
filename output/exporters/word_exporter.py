import logging
import os

from docx import Document
from docx.shared import Pt

from output.formatters.document_formatter import CONFIDENTIAL_PLACEHOLDER, InterimDocument

logger = logging.getLogger(__name__)


def generate_docx(document: InterimDocument, output_path: str) -> str:
    """Write a .docx file from an InterimDocument.

    Returns the absolute path of the written file.
    """
    doc = Document()

    # Document properties
    doc.core_properties.author = "KnowledgeKeeper"
    doc.core_properties.title = document.title

    # Title page
    title_para = doc.add_heading(document.title, level=0)
    doc.add_paragraph(document.subtitle)
    doc.add_paragraph("")  # spacer

    if document.confidential_sections_note:
        notice = doc.add_paragraph()
        run = notice.add_run(
            f"Confidentiality Notice: Some sections of this document have been redacted. "
            f"Flagged areas: {document.confidential_sections_note}"
        )
        run.italic = True
        doc.add_paragraph("")

    # Sections
    for section in document.sections:
        if section.is_confidential:
            heading = doc.add_heading(f"{section.heading} [CONFIDENTIAL]", level=1)
            doc.add_paragraph(CONFIDENTIAL_PLACEHOLDER)
        else:
            doc.add_heading(section.heading, level=1)
            _render_markdown_to_docx(doc, section.content_markdown)

    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    logger.info("session=%s docx written to %s", document.session_id, abs_path)
    return abs_path


def _render_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    """Convert basic markdown to python-docx paragraphs."""
    lines = markdown_text.split("\n")
    in_list = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if in_list:
                in_list = False
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("### "):
            in_list = False
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            in_list = False
            doc.add_heading(stripped[3:], level=2)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            in_list = True
            doc.add_paragraph(stripped[2:], style="List Bullet")

        # Gap markers
        elif stripped.startswith("[GAP:"):
            in_list = False
            para = doc.add_paragraph()
            run = para.add_run("KNOWLEDGE GAP: ")
            run.bold = True
            # Extract description from [GAP: description]
            gap_text = stripped.strip("[]")
            if gap_text.startswith("GAP:"):
                gap_text = gap_text[4:].strip()
            para.add_run(gap_text)

        # Bold lines (** wrapped)
        elif stripped.startswith("**") and stripped.endswith("**"):
            in_list = False
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip("*").strip())
            run.bold = True

        # Normal text
        else:
            in_list = False
            doc.add_paragraph(stripped)
