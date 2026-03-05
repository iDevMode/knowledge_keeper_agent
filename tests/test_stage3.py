import json
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from langchain_core.messages import AIMessage, HumanMessage

from agents.stage3_document_generation.generator import (
    GenerationRequest,
    GenerationResult,
    _validate_output,
    generate_document,
)
from agents.stage3_document_generation.prompts import (
    build_context_block,
    _format_answers_by_block,
    _format_profile_for_context,
    _format_risk_flags,
)
from models.risk_flags import RiskFlag
from models.role_intelligence_profile import RoleIntelligenceProfile
from output.formatters.document_formatter import (
    CONFIDENTIAL_PLACEHOLDER,
    InterimDocument,
    DocumentSection,
    parse_llm_output,
    _apply_confidentiality_filter,
)
from output.exporters.word_exporter import generate_docx
from output.exporters.pdf_exporter import _build_html, _markdown_to_html


# ---- Helpers ----

FIXTURES_PATH = Path(__file__).parent / "fixtures"
PROFILES_PATH = FIXTURES_PATH / "sample_role_profiles.json"
STAGE2_PATH = FIXTURES_PATH / "sample_stage2_results.json"


def _load_profile(profile_id: str) -> RoleIntelligenceProfile:
    with open(PROFILES_PATH) as f:
        fixtures = json.load(f)
    return RoleIntelligenceProfile.model_validate(fixtures[profile_id])


def _load_stage2_result(result_id: str) -> dict:
    with open(STAGE2_PATH) as f:
        fixtures = json.load(f)
    return fixtures[result_id]


def _reconstruct_conversation_history(raw: list):
    messages = []
    for entry in raw:
        if entry["type"] == "ai":
            messages.append(AIMessage(content=entry["content"]))
        else:
            messages.append(HumanMessage(content=entry["content"]))
    return messages


def _make_generation_request(result_id: str = "process_heavy") -> GenerationRequest:
    result = _load_stage2_result(result_id)
    profile = _load_profile(result["profile_id"])
    history = _reconstruct_conversation_history(result["conversation_history"])
    risk_flags = [RiskFlag.model_validate(rf) for rf in result["risk_flags"]]

    return GenerationRequest(
        session_id=result["session_id"],
        profile=profile,
        conversation_history=history,
        risk_flags=risk_flags,
        answers=result["answers"],
        block_order=result["block_order"],
        block_depths=result["block_depths"],
    )


def _build_valid_markdown(profile_id: str = "process_heavy") -> str:
    """Build synthetic LLM output with all required sentinels."""
    profile = _load_profile(profile_id)
    data = profile.model_dump()
    p1 = str(data.get("priority_1", "Priority 1")).replace("_", " ").title()
    p2 = str(data.get("priority_2", "Priority 2")).replace("_", " ").title()
    p3 = str(data.get("priority_3", "Priority 3")).replace("_", " ").title()

    return f"""\
### SECTION: Document Header

**Handover Document**
Role: {data.get('job_title', 'Unknown')}
Company: {data.get('company_name', 'Unknown')}
Date: 28 February 2026
Prepared by: KnowledgeKeeper | Nukode

### SECTION: Risk Summary

## Critical
- SAP scheduling workaround is a single point of failure

## High
- Supplier agreements are informal

### SECTION: Role Overview

This role is the operational backbone of the production team. The Senior Operations Coordinator manages weekly production scheduling, supplier coordination, and reporting.

### SECTION: Knowledge Transfer — {p1}

## What Was Captured
The weekly production scheduling process was documented in full detail.

## Known Gaps
[GAP: Detailed SAP custom report configuration steps were not fully captured]

### SECTION: Knowledge Transfer — {p2}

## What Was Captured
Full system inventory and access details were captured.

### SECTION: Knowledge Transfer — {p3}

## What Was Captured
Undocumented workarounds were captured comprehensively.

### SECTION: Key Relationships

- AlphaSteel: Sarah Chen, Account Manager. Priority scheduling arrangement.
- BetaChem: Dave Peters. Informal 12% discount.
- GammaPlastics: Rocky relationship, missed deliveries.

### SECTION: Systems and Access

- SAP ERP: Production module admin, custom report ZPPROD001
- Power BI: Workspace owner, built dashboards from scratch
- Excel: Three custom macros for backup scheduling

### SECTION: In-Flight Items

- AlphaSteel factory move in Q3 — potential 2-3 week supply disruption
- BetaChem discount formalisation needed

### SECTION: Decision-Making Logic and Judgment

**Scenario: Batch Split vs Single Run**
Context: Weekly scheduling decision
Approach: Weigh shift team capability, material availability, delivery deadlines
Rule of thumb: Never sacrifice quality for efficiency

### SECTION: Undocumented Knowledge

**SAP Batch Scheduling Workaround**
What it is: Custom 4-hour Monday process replacing the standard 2-hour broken run
Why it matters: Standard SAP output gives wrong results for specialty products
How to apply it: Export data, run through Excel macro, re-import

### SECTION: Advice to Your Replacement

The most important thing to understand is that the SAP system documentation is outdated. Trust what works, not what the manual says. Meet the shift supervisors in your first week. Never promise a delivery date without checking the SAP schedule first.

### SECTION: Recommended Onboarding Plan

**Week 1 — Orientation**
Shadow the Monday scheduling run. Meet all three shift supervisors.

**Week 2 — Relationship Building**
Call AlphaSteel (Sarah Chen) and BetaChem (Dave Peters).

**Week 3 — Process Immersion**
Run the scheduling process with oversight.

**Week 4 — Independence**
Run the full scheduling cycle independently.

### SECTION: Knowledge Gaps and Recommended Follow-Up

- [GAP: Complete SAP custom report configuration] — Risk: High. Ask the SAP consultant for report ZPPROD001 documentation.
- [GAP: Excel macro source code review] — Risk: Medium. Macros are on the departing employee's desktop in 'Ops Tools' folder.

This document has been generated by KnowledgeKeeper. It is based on 30 questions across two interview sessions. 4 risk flags were identified and are summarised in Section 2.
"""


# ---- TestContextAssembly ----

class TestContextAssembly:
    def test_profile_context_included(self):
        req = _make_generation_request("process_heavy")
        context = build_context_block(
            req.profile, req.conversation_history, req.risk_flags,
            req.answers, req.block_order, req.block_depths,
        )
        assert "Senior Operations Coordinator" in context
        assert "Precision Manufacturing" in context

    def test_risk_flags_formatted_by_severity(self):
        req = _make_generation_request("process_heavy")
        formatted = _format_risk_flags(req.risk_flags)
        # Critical should appear before high
        critical_pos = formatted.find("CRITICAL")
        high_pos = formatted.find("HIGH")
        assert critical_pos < high_pos

    def test_answers_grouped_by_block(self):
        req = _make_generation_request("process_heavy")
        formatted = _format_answers_by_block(req.answers, req.block_order, req.block_depths)
        # Role orientation should appear first
        orientation_pos = formatted.find("Role Orientation")
        block_pos = formatted.find("Internal Processes Workflows")
        assert orientation_pos < block_pos

    def test_no_risk_flags_omits_section(self):
        req = _make_generation_request("process_heavy")
        context = build_context_block(
            req.profile, req.conversation_history, [],  # empty risk flags
            req.answers, req.block_order, req.block_depths,
        )
        assert "RISK FLAGS IDENTIFIED" not in context

    def test_confidential_warning_injected(self):
        req = _make_generation_request("process_heavy")
        # process_heavy has confidential_sections="Performance-related notes"
        context = build_context_block(
            req.profile, req.conversation_history, req.risk_flags,
            req.answers, req.block_order, req.block_depths,
        )
        assert "Performance-related notes" in context


# ---- TestDocumentGenerator ----

class TestDocumentGenerator:
    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_generate_document_success(self, mock_get_llm):
        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content=_build_valid_markdown())
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("process_heavy")
        result = generate_document(req)

        assert isinstance(result, GenerationResult)
        assert result.session_id == req.session_id
        assert len(result.raw_markdown) > 0
        assert result.generation_metadata["question_count"] == len(req.answers)

    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_generate_document_retries_on_missing_sections(self, mock_get_llm):
        # First call missing Risk Summary, second call complete
        incomplete = _build_valid_markdown().replace("### SECTION: Risk Summary", "### SECTION: Risk Overview")
        complete = _build_valid_markdown()

        mock_llm = MagicMock()
        mock_llm.invoke.side_effect = [
            MagicMock(content=incomplete),
            MagicMock(content=complete),
        ]
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("process_heavy")
        result = generate_document(req)

        assert mock_llm.invoke.call_count == 2
        assert "Risk Summary" in result.raw_markdown

    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_generate_document_raises_after_two_failures(self, mock_get_llm):
        incomplete = "Just some random text without any sentinel markers."

        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content=incomplete)
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("process_heavy")
        with pytest.raises(RuntimeError, match="missing sections"):
            generate_document(req)

        assert mock_llm.invoke.call_count == 2

    @patch("agents.stage3_document_generation.generator.ChatAnthropic")
    def test_max_tokens_is_8192(self, mock_cls):
        from agents.stage3_document_generation.generator import _get_generation_llm
        _get_generation_llm()
        mock_cls.assert_called_once()
        call_kwargs = mock_cls.call_args
        assert call_kwargs.kwargs.get("max_tokens") == 8192


# ---- TestOutputValidation ----

class TestOutputValidation:
    def test_valid_output_passes(self):
        markdown = _build_valid_markdown()
        missing = _validate_output(markdown)
        assert missing == []

    def test_missing_section_detected(self):
        markdown = _build_valid_markdown().replace("### SECTION: Risk Summary", "")
        missing = _validate_output(markdown)
        assert "Risk Summary" in missing

    def test_insufficient_knowledge_transfer_detected(self):
        # Remove two of the three Knowledge Transfer sections
        markdown = _build_valid_markdown()
        parts = markdown.split("### SECTION: Knowledge Transfer")
        # Keep only the first Knowledge Transfer, remove the other two
        if len(parts) >= 4:
            markdown = parts[0] + "### SECTION: Knowledge Transfer" + parts[1] + "### SECTION: Key Relationships" + parts[3].split("### SECTION: Key Relationships", 1)[-1]
            missing = _validate_output(markdown)
            assert any("Knowledge Transfer" in m for m in missing)


# ---- TestDocumentFormatter ----

class TestDocumentFormatter:
    def test_parses_all_sections(self):
        raw = _build_valid_markdown()
        profile = _load_profile("process_heavy")
        doc = parse_llm_output(raw, profile, "test-session")
        assert len(doc.sections) == 14

    def test_section_names_correct(self):
        raw = _build_valid_markdown()
        profile = _load_profile("process_heavy")
        doc = parse_llm_output(raw, profile, "test-session")
        assert doc.sections[0].name == "Document Header"
        assert doc.sections[1].name == "Risk Summary"

    def test_confidential_sections_marked(self):
        raw = _build_valid_markdown()
        profile = _load_profile("process_heavy")
        # process_heavy has confidential_sections="Performance-related notes"
        doc = parse_llm_output(raw, profile, "test-session")
        # No section should match "performance" in its name for this document
        # Test with a custom profile that has matching confidential sections
        sections = [
            DocumentSection(name="Key Relationships", heading="Key Relationships",
                           content_markdown="Some client notes here", section_number=1),
            DocumentSection(name="Systems and Access", heading="Systems and Access",
                           content_markdown="Technical details", section_number=2),
        ]
        result = _apply_confidentiality_filter(sections, "client notes")
        assert result[0].is_confidential is True
        assert result[1].is_confidential is False

    def test_confidential_content_redacted(self):
        sections = [
            DocumentSection(name="Client Details", heading="Client Details",
                           content_markdown="Sensitive client financial information here",
                           section_number=1),
        ]
        result = _apply_confidentiality_filter(sections, "client")
        assert result[0].content_markdown == CONFIDENTIAL_PLACEHOLDER

    def test_no_confidential_sections(self):
        raw = _build_valid_markdown("relationship_heavy")
        profile = _load_profile("relationship_heavy")
        # relationship_heavy has confidential_sections=None
        doc = parse_llm_output(raw, profile, "test-session")
        for section in doc.sections:
            assert section.is_confidential is False

    def test_gap_markers_preserved(self):
        raw = _build_valid_markdown()
        profile = _load_profile("process_heavy")
        doc = parse_llm_output(raw, profile, "test-session")
        # Find the Knowledge Transfer section which contains a GAP marker
        kt_sections = [s for s in doc.sections if "Knowledge Transfer" in s.name]
        assert any("[GAP:" in s.content_markdown for s in kt_sections)

    def test_document_title_includes_job_title(self):
        raw = _build_valid_markdown()
        profile = _load_profile("process_heavy")
        doc = parse_llm_output(raw, profile, "test-session")
        assert "Senior Operations Coordinator" in doc.title


# ---- TestWordExporter ----

class TestWordExporter:
    def test_generates_docx_file(self):
        doc = InterimDocument(
            session_id="test",
            title="Handover Document — Test Role",
            subtitle="Test Company | Generated 28 February 2026",
            sections=[
                DocumentSection(name="Role Overview", heading="Role Overview",
                               content_markdown="This is the role overview.", section_number=1),
            ],
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "test_output.docx")
            result_path = generate_docx(doc, output_path)
            assert os.path.exists(result_path)
            assert result_path.endswith(".docx")

    def test_docx_contains_title(self):
        from docx import Document as DocxDocument

        doc = InterimDocument(
            session_id="test",
            title="Handover Document — Senior Ops Coordinator",
            subtitle="Test Co",
            sections=[
                DocumentSection(name="Overview", heading="Overview",
                               content_markdown="Test content.", section_number=1),
            ],
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "test.docx")
            generate_docx(doc, output_path)
            docx_doc = DocxDocument(output_path)
            all_text = " ".join(p.text for p in docx_doc.paragraphs)
            assert "Senior Ops Coordinator" in all_text

    def test_confidential_sections_redacted_in_docx(self):
        from docx import Document as DocxDocument

        doc = InterimDocument(
            session_id="test",
            title="Test",
            subtitle="Test",
            sections=[
                DocumentSection(name="Secret Section", heading="Secret Section",
                               content_markdown=CONFIDENTIAL_PLACEHOLDER,
                               is_confidential=True, section_number=1),
            ],
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "test.docx")
            generate_docx(doc, output_path)
            docx_doc = DocxDocument(output_path)
            all_text = " ".join(p.text for p in docx_doc.paragraphs)
            assert "[CONFIDENTIAL]" in all_text
            assert CONFIDENTIAL_PLACEHOLDER in all_text

    def test_heading_styles_applied(self):
        from docx import Document as DocxDocument

        doc = InterimDocument(
            session_id="test",
            title="Test",
            subtitle="Test",
            sections=[
                DocumentSection(name="Test Section", heading="Test Section",
                               content_markdown="## Sub Heading\n\nSome content.",
                               section_number=1),
            ],
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "test.docx")
            generate_docx(doc, output_path)
            docx_doc = DocxDocument(output_path)
            heading_styles = [p.style.name for p in docx_doc.paragraphs if "Heading" in (p.style.name or "")]
            assert any("Heading 2" in s for s in heading_styles)


# ---- TestPdfExporter ----

class TestPdfExporter:
    def test_raises_if_weasyprint_unavailable(self):
        with patch("output.exporters.pdf_exporter.WEASYPRINT_AVAILABLE", False):
            from output.exporters.pdf_exporter import generate_pdf
            doc = InterimDocument(
                session_id="test", title="Test", subtitle="Test", sections=[],
            )
            with pytest.raises(RuntimeError, match="WeasyPrint"):
                generate_pdf(doc, "test.pdf")

    @patch("output.exporters.pdf_exporter.WEASYPRINT_AVAILABLE", True)
    @patch("output.exporters.pdf_exporter.WeasyHTML")
    def test_generates_pdf_file(self, mock_weasy):
        from output.exporters.pdf_exporter import generate_pdf

        doc = InterimDocument(
            session_id="test",
            title="Test Document",
            subtitle="Test",
            sections=[
                DocumentSection(name="Overview", heading="Overview",
                               content_markdown="Content here.", section_number=1),
            ],
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "test.pdf")
            generate_pdf(doc, output_path)
            mock_weasy.assert_called_once()

    def test_html_contains_title(self):
        doc = InterimDocument(
            session_id="test",
            title="Handover — Senior Coordinator",
            subtitle="Test Co",
            sections=[
                DocumentSection(name="Overview", heading="Overview",
                               content_markdown="Content.", section_number=1),
            ],
        )
        html = _build_html(doc)
        assert "Senior Coordinator" in html

    def test_confidential_sections_have_css_class(self):
        doc = InterimDocument(
            session_id="test",
            title="Test",
            subtitle="Test",
            sections=[
                DocumentSection(name="Secret", heading="Secret",
                               content_markdown=CONFIDENTIAL_PLACEHOLDER,
                               is_confidential=True, section_number=1),
            ],
        )
        html = _build_html(doc)
        assert "class='confidential'" in html


# ---- TestHtmlBuilder ----

class TestHtmlBuilder:
    def test_headings_converted(self):
        html = _markdown_to_html("## Sub Heading\n\nContent here")
        assert "<h2>" in html
        assert "Sub Heading" in html

    def test_bullets_converted(self):
        html = _markdown_to_html("- Item one\n- Item two")
        assert "<li>" in html
        assert "<ul>" in html

    def test_gap_markers_styled(self):
        html = _markdown_to_html("[GAP: Missing SAP documentation]")
        assert "class='gap'" in html
        assert "KNOWLEDGE GAP" in html

    def test_plain_text_wrapped(self):
        html = _markdown_to_html("Just a normal paragraph.")
        assert "<p>" in html


# ---- TestIntegration ----

class TestIntegration:
    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_process_heavy_end_to_end(self, mock_get_llm):
        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content=_build_valid_markdown("process_heavy"))
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("process_heavy")
        result = generate_document(req)
        doc = parse_llm_output(result.raw_markdown, req.profile, req.session_id)

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = os.path.join(tmp_dir, "process_heavy.docx")
            generate_docx(doc, output_path)
            assert os.path.exists(output_path)
            assert doc.sections[0].name == "Document Header"
            assert doc.has_risk_flags is True

    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_decision_heavy_end_to_end(self, mock_get_llm):
        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content=_build_valid_markdown("decision_heavy"))
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("decision_heavy")
        result = generate_document(req)
        doc = parse_llm_output(result.raw_markdown, req.profile, req.session_id)

        assert len(doc.sections) == 14
        assert "Senior Financial Advisor" in doc.title

    @patch("agents.stage3_document_generation.generator._get_generation_llm")
    def test_relationship_heavy_end_to_end(self, mock_get_llm):
        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content=_build_valid_markdown("relationship_heavy"))
        mock_get_llm.return_value = mock_llm

        req = _make_generation_request("relationship_heavy")
        result = generate_document(req)
        doc = parse_llm_output(result.raw_markdown, req.profile, req.session_id)

        assert len(doc.sections) == 14
        assert "Principal Consultant" in doc.title
        # relationship_heavy has no confidential sections
        for section in doc.sections:
            assert section.is_confidential is False
